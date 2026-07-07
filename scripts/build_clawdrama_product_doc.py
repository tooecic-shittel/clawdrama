from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/djj/Documents/爪爪短剧")
OUT = ROOT / "docs" / "ClawDrama短剧工作流产品介绍.docx"
SHOT_DIR = ROOT / "docs_assets" / "clawdrama_workflow_screens"
BRAND_DIR = ROOT / "docs_assets" / "brand"

FONT = "Hiragino Sans GB"
TITLE_FONT = "Hiragino Sans GB"

DARK = RGBColor(16, 24, 40)
INK = RGBColor(45, 55, 72)
MUTED = RGBColor(101, 116, 139)
GREEN = RGBColor(118, 214, 38)
GREEN_DARK = RGBColor(61, 132, 14)
BLUE = RGBColor(37, 99, 235)
LINE = "D8E0EA"
SOFT_GREEN = "F3FCEB"
SOFT_BLUE = "EFF6FF"
SOFT_GRAY = "F6F8FB"
PALE = "FAFBFD"


def set_run_font(run, size=None, color=None, bold=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.bold = bold


def add_run(p, text, size=None, color=None, bold=False):
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), sz)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top=110, start=145, bottom=110, end=145):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_col_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    add_run(p, text, color=BLUE if level == 1 else DARK, bold=True, size=16 if level == 1 else 11.5)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, color=DARK, bold=True, size=10.5)
        add_run(p, text[len(bold_prefix):], color=INK, size=10.5)
    else:
        add_run(p, text, color=INK, size=10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        add_run(p, item, color=INK, size=9.7)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, text, color=MUTED, size=8.2)


def add_image(doc, filename, caption, width=6.55):
    p = doc.add_paragraph()
    keep_with_next(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    path = SHOT_DIR / filename
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_highlight(doc, label, text, fill=SOFT_GREEN, accent=GREEN_DARK, after_space=True):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_col_width(cell, 6.45)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "D7E7D0" if fill == SOFT_GREEN else LINE, "6")
    set_cell_margins(cell, top=135, bottom=135, start=170, end=170)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, label, color=accent, bold=True, size=10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    add_run(p2, text, color=INK, size=9.6)
    if after_space:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_metric_cards(doc):
    rows = [
        ("全流程", "从剧本到成片导出闭环"),
        ("镜头级", "每个分镜可独立生成与重试"),
        ("资产化", "角色、场景、首尾帧复用"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (top, bottom) in enumerate(rows):
        cell = table.cell(0, i)
        set_col_width(cell, 2.08)
        set_cell_shading(cell, PALE)
        set_cell_border(cell, LINE, "6")
        set_cell_margins(cell, top=140, bottom=125, start=130, end=130)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, top, color=GREEN_DARK, bold=True, size=13)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        add_run(p2, bottom, color=MUTED, size=8.8)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_col_width(cell, widths[i])
        set_cell_shading(cell, "EAF1F8")
        set_cell_border(cell, "C9D6E5")
        set_cell_margins(cell, top=115, bottom=115, start=125, end=125)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        add_run(p, header, color=DARK, bold=True, size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_col_width(cell, widths[i])
            set_cell_shading(cell, "FFFFFF")
            set_cell_border(cell)
            set_cell_margins(cell, top=100, bottom=100, start=125, end=125)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.08
            add_run(p, str(text), color=INK, size=8.7)
    doc.add_paragraph()
    return table


def setup_styles(doc):
    styles = doc.styles
    for name in ("Normal", "List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), FONT)
        style.font.size = Pt(10.2 if name == "Normal" else 9.7)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(5 if name == "Normal" else 3)
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 15, 6),
        ("Heading 2", 12.2, DARK, 9, 4),
        ("Heading 3", 10.8, DARK, 8, 3),
    ):
        style = styles[name]
        style.font.name = TITLE_FONT
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), TITLE_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def setup_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.66)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "ClawDrama 产品工作流介绍", color=MUTED, size=8)


def add_cover(doc):
    logo = BRAND_DIR / "claw-logo-lockup.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_picture(str(logo), width=Inches(1.55))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "AI 短剧全流程生产工作流产品介绍", color=DARK, bold=True, size=24)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_run(p, "从剧本输入、AI 改写、角色场景提取、分镜拆解，到图片、视频、配音、合成与整集导出的一体化生产平台。", color=MUTED, size=10.5)

    add_metric_cards(doc)
    add_image(doc, "00_website_home_cover.png", "图 1：ClawDrama 网站首页与产品定位", width=6.55)
    add_highlight(
        doc,
        "核心定位",
        "ClawDrama 不是单点 AI 工具，而是面向短剧生产的工作流系统。它把创意文本拆解成可编辑、可复用、可计费、可规模化交付的生产节点。",
        fill=SOFT_GREEN,
    )
    doc.add_page_break()


def workflow_overview(doc):
    add_heading(doc, "1. 产品概览", 1)
    add_body(doc, "爪爪短剧 ClawDrama 围绕短剧团队真实生产链路设计，将剧本、角色、场景、分镜、图片、视频、配音和导出统一组织在项目工作台内。团队可以顺着流程自动推进，也可以在关键节点人工编辑、局部重抽或替换素材。")
    add_body(doc, "平台采用“项目 — 剧集 — 分镜 — 素材 — 成片”的结构，把抽象创作过程拆成清晰的生产单元，适合内容团队协作、项目复盘、成本核算和后续规模化运营。")
    add_highlight(doc, "Highlight", "用 AI 处理高频、结构化、重复性的生产任务，把人的判断留给故事节奏、角色表现、镜头审美和商业决策；每一步都有明确输入、处理逻辑和产出物，既能自动化推进，又保留人工修正空间。", fill=SOFT_BLUE, accent=BLUE)

    add_heading(doc, "2. 端到端工作流总览", 1)
    rows = [
        ["01", "项目入口", "建立项目、进入剧集、查看制作状态", "项目工作台"],
        ["02", "剧本处理", "导入原文，AI 改写为短剧化脚本", "可生产剧本"],
        ["03", "资产提取", "抽取角色、场景、性别、声音策略", "角色库 / 场景库"],
        ["04", "分镜拆解", "拆成镜头，生成画面和视频提示词", "分镜表"],
        ["05", "视觉生成", "生成角色图、场景图、镜头首尾帧", "视觉资产"],
        ["06", "视频生产", "图生视频、配音、音画合成", "镜头视频"],
        ["07", "成片导出", "整集拼接、预览、下载或精剪", "交付成片"],
    ]
    add_table(doc, ["阶段", "模块", "核心动作", "关键产出"], rows, [0.6, 1.15, 3.15, 1.55])
    doc.add_page_break()


STAGES = [
    {
        "title": "项目工作台与剧集入口",
        "image": "01_projects_dashboard.png",
        "goal": "以项目为单位管理短剧内容，快速查看剧集数量、制作状态、更新时间和进入入口。",
        "capabilities": ["项目级组织：剧本、角色、场景、分镜和素材统一挂载。", "多项目并行：支持工作室同时推进多个短剧或客户项目。", "状态可视化：用进度和状态减少团队沟通成本。"],
        "outputs": "项目列表、剧集入口、制作进度概览",
        "value": "把创作灵感变成可管理的生产项目，解决多短剧、多剧集、多素材分散的问题。",
    },
    {
        "title": "原始内容输入",
        "image": "02_original_content.png",
        "goal": "接收小说原文、剧情梗概、已有剧本或创意片段，为后续 AI 改写和结构化处理提供素材基础。",
        "capabilities": ["支持非标准输入：不要求用户先准备专业剧本格式。", "原文留痕：便于对比 AI 改写前后的剧情表达。", "流程衔接：原始内容可直接进入改写、提取和分镜。"],
        "outputs": "原始故事文本、可追溯素材源、后续 AI 处理上下文",
        "value": "降低创作入口门槛，让小说、梗概、口述灵感都能进入短剧生产链。",
    },
    {
        "title": "AI 剧本改写",
        "image": "03_ai_rewrite.png",
        "goal": "将原始文本转换为更适合短剧视频生产的格式化剧本，强化开场钩子、冲突推进、对白节奏和视觉可拍性。",
        "capabilities": ["小说转剧本：把叙事性文字转为短剧拍摄脚本。", "节奏优化：突出强开场、冲突、反转和结尾悬念。", "格式整理：为角色提取和分镜拆解提供稳定结构。"],
        "outputs": "格式化剧本、短剧化对白与动作、适合分镜拆解的文本结构",
        "value": "把“可读的故事”升级为“可生产的视频脚本”，减少编剧初稿整理时间。",
    },
    {
        "title": "角色与场景提取",
        "image": "04_extract.png",
        "goal": "从剧本中自动识别人物、场景和基础设定，形成图片生成、视频参考和音色分配所需的资产库。",
        "capabilities": ["角色识别：抽取名称、身份、外貌、性格和剧情定位。", "场景识别：抽取空间特征、时间氛围和视觉提示。", "去重合并：降低同一角色或场景重复创建的概率。"],
        "outputs": "角色库、场景库、角色/场景生成提示词",
        "value": "建立统一资产底座，让后续镜头不再孤立生成，提高人物和空间一致性。",
    },
    {
        "title": "角色性别与旁白音色设定",
        "image": "05_gender_voice.png",
        "goal": "为角色建立基础声音策略，区分角色对白、旁白和内心独白的表现方式。",
        "capabilities": ["性别辅助判断：根据角色名称、描述和对白语境推断。", "旁白音色管理：将旁白/解说音频与角色设定关联。", "人工兜底：用户可随时调整不符合角色定位的设定。"],
        "outputs": "角色性别设定、旁白/独白音色策略、配音生成配置",
        "value": "让声音生产提前结构化，减少后续配音错位、性别错配和角色表现不一致。",
    },
    {
        "title": "AI 分镜拆解",
        "image": "06_storyboard.png",
        "goal": "将剧本拆解成可生产的视频镜头单元，每个镜头包含画面、角色、场景、对白、时长和视频提示词。",
        "capabilities": ["镜头结构化：把长文本拆为可管理的视频生产单元。", "提示词生成：为图像和视频模型准备具体视觉描述。", "分镜级编辑：每个镜头都可以独立修改和重试。"],
        "outputs": "分镜列表、镜头画面描述、视频生成提示词、角色/场景绑定",
        "value": "把剧本从“文本资产”转化为“镜头生产计划”，是 AI 短剧规模化生产的关键。",
    },
    {
        "title": "角色形象生成",
        "image": "07_character_images.png",
        "goal": "根据角色设定生成可复用的人物视觉资产，用于镜头图片和视频参考，提升人物一致性。",
        "capabilities": ["角色图生成：将文字设定转为可视化人物形象。", "视觉一致性：同一角色图可在多个镜头中复用。", "多视图支持：可扩展侧面、背面等参考。"],
        "outputs": "角色正面图、角色参考图、可复用人物资产",
        "value": "减少人物反复描述成本，让短剧角色在不同镜头中保持识别度。",
    },
    {
        "title": "场景图片生成",
        "image": "08_scene_images.png",
        "goal": "根据场景库生成环境视觉资产，为镜头构图、空间氛围和视频参考提供统一基底。",
        "capabilities": ["空间视觉化：将办公室、街道、宫殿等场景转为图片。", "风格统一：保持项目色调、时代感和美术方向。", "素材复用：同一场景可支撑多个镜头生成。"],
        "outputs": "场景图、环境参考图、项目视觉基调",
        "value": "让短剧拥有稳定世界观和空间连续性，而不是每个镜头重新随机生成。",
    },
    {
        "title": "镜头图片与首尾帧生成",
        "image": "09_shot_images.png",
        "goal": "为每个分镜生成首帧、尾帧或分镜图，作为视频生成的核心视觉约束。",
        "capabilities": ["首帧控制：决定视频起始画面。", "尾帧控制：帮助视频形成明确动作变化和结尾画面。", "分镜图资产化：每个镜头拥有可预览的视觉草图。"],
        "outputs": "镜头首帧、镜头尾帧、分镜图、视频参考图",
        "value": "把抽象镜头描述变成可见画面，大幅提升视频生成可控性。",
    },
    {
        "title": "图生视频生成",
        "image": "10_video_generation.png",
        "goal": "基于镜头提示词、首尾帧和参考图生成单镜头视频，支持按镜头重试和批量生产。",
        "capabilities": ["多引擎适配：支持不同视频模型链路。", "参考图生视频：通过角色图、场景图、首尾帧约束画面。", "分镜级容错：失败镜头可单独重试，不影响其他镜头。"],
        "outputs": "单镜头视频、模型生成记录、可供合成/剪辑的视频素材",
        "value": "将短剧生产从“整片一次生成”改为“镜头级稳定生产”，更接近专业影视制作逻辑。",
    },
    {
        "title": "旁白配音与音频生成",
        "image": "11_voiceover.png",
        "goal": "为旁白、内心独白或解说类文本生成音频，补足视频模型无法稳定表达的声音部分。",
        "capabilities": ["TTS 生成：将文字转换为自然语音。", "角色音色关联：旁白和独白可以保持声音一致。", "音频资产管理：生成结果进入后续合成流程。"],
        "outputs": "旁白音频、独白音频、镜头音频素材",
        "value": "让短剧具备完整的信息传达能力，尤其适合小说改编、内心戏和解说型短剧。",
    },
    {
        "title": "视频合成",
        "image": "12_composition.png",
        "goal": "将单镜头视频与音频、字幕等素材整合为可连续播放的镜头片段。",
        "capabilities": ["音画同步：将视频与 TTS 音频按时间轴组合。", "字幕处理：为内容理解和平台发布提供基础字幕能力。", "片段标准化：让每个镜头具备统一编码和导出格式。"],
        "outputs": "合成镜头片段、字幕视频、可拼接素材",
        "value": "把 AI 分散生成的素材变成可剪辑、可拼接、可交付的视频片段。",
    },
    {
        "title": "整集拼接与导出",
        "image": "13_export.png",
        "goal": "将所有镜头按顺序拼接为完整短剧成片，并提供下载和后续精剪入口。",
        "capabilities": ["镜头排序拼接：按分镜顺序组成完整剧集。", "成片预览：快速判断整集完整度。", "后期衔接：保留专业剪辑软件继续精修的空间。"],
        "outputs": "整集视频、导出文件、可进入发布或精剪的成片",
        "value": "完成从文本到成片的闭环，让 AI 生产真正进入可交付状态。",
    },
]


def add_stage(doc, idx, info):
    add_heading(doc, f"{idx + 3}. {info['title']}", 1)
    add_image(doc, info["image"], f"图 {idx + 2}：{info['title']}界面截图", width=6.55)
    add_highlight(doc, "阶段定位", info["goal"], fill=SOFT_BLUE, accent=BLUE)
    add_heading(doc, "关键能力", 3)
    add_bullets(doc, info["capabilities"])
    add_highlight(doc, "产出亮点", info["outputs"], fill=SOFT_GREEN, accent=GREEN_DARK)
    add_body(doc, f"业务价值：{info['value']}", bold_prefix="业务价值：")


def summary_sections(doc):
    doc.add_page_break()
    add_heading(doc, "16. 产品优势总结", 1)
    add_table(
        doc,
        ["优势", "说明", "客户价值"],
        [
            ["全流程一体化", "剧本、角色、场景、分镜、图片、视频、配音和导出集中在一个工作台。", "减少跨工具复制、上传、命名和对齐成本。"],
            ["分镜级可控", "每个镜头都能独立编辑、生成、重试和替换。", "比“一键生成整片”更适合真实短剧生产。"],
            ["资产可复用", "角色图、场景图和首尾帧可作为跨镜头参考。", "提升人物、空间和视觉风格一致性。"],
            ["模型可扩展", "不同任务可接入不同图片、视频和语音模型。", "在质量、速度、成本和风格之间灵活平衡。"],
            ["运营可计费", "具备积分、余额、流水和套餐等商业化基础。", "适合团队管理和 SaaS 化运营。"],
        ],
        [1.35, 3.0, 2.1],
    )

    add_heading(doc, "17. 适用场景与目标用户", 1)
    add_bullets(
        doc,
        [
            "短剧团队：快速完成剧本视觉化、分镜生产和样片生成。",
            "小说 IP 团队：将小说原文或大纲转换为短剧剧集和镜头资产。",
            "MCN 与账号运营：批量测试不同剧情方向和投放内容。",
            "品牌营销团队：用剧情化方式包装产品卖点，生成短视频广告素材。",
            "个人创作者：降低短剧制作门槛，以较低成本完成从创意到成片的尝试。",
        ],
    )

    add_heading(doc, "18. 结语", 1)
    add_body(doc, "ClawDrama 的本质是一套 AI 辅助的短剧生产基础设施。它不追求用一次生成替代所有创作判断，而是把短剧制作拆解成专业、可控、可复用的生产节点，让团队能持续迭代故事、角色、画面和成片质量。")
    add_body(doc, "对于创作者，它降低了从文字到视频的门槛；对于内容团队，它提升了前期开发和素材生产效率；对于运营方，它提供了可计费、可扩展、可规模化的 AI 内容生产平台。")
    add_highlight(doc, "最终价值", "让短剧从“灵感驱动的手工作坊”走向“节点清晰、资产沉淀、质量可控、成本可算”的 AI 内容生产线。", after_space=False)


def build():
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    add_cover(doc)
    workflow_overview(doc)
    for idx, info in enumerate(STAGES, start=1):
        add_stage(doc, idx, info)
    summary_sections(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
